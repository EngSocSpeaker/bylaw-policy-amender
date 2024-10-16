import json
import re
from typing import cast
from PySide6.QtCore import (
    QAbstractItemModel, QAbstractTableModel, QModelIndex,
    QPersistentModelIndex, Qt
)
from PySide6.QtWidgets import (
    QStyledItemDelegate, QComboBox, QWidget, QStyleOptionViewItem, QLineEdit,
    QTextEdit
)
import docx
import docx.shared

from github import gh
from str_manip import TeXSource, SectionValidator

ROLES = {
    Qt.ItemDataRole.DisplayRole,
    Qt.ItemDataRole.EditRole,
    Qt.ItemDataRole.AccessibleTextRole,
}
COLUMNS = ['File', 'Section', 'Current text', 'Proposed text']

class TreeFileDelegate(QStyledItemDelegate):
    def createEditor(self, parent: QWidget, option: QStyleOptionViewItem, index: QModelIndex | QPersistentModelIndex) -> QWidget:
        cb = QComboBox(parent)
        cb.addItems([item['path'] for item in gh.getTree('skule/bylaws') if item['path'].endswith('.tex')])
        self.setEditorData(cb, index)
        cb.currentIndexChanged.connect(lambda: self.commitData.emit(cb))
        return cb

    def setEditorData(self, cb: QComboBox, index: QModelIndex | QPersistentModelIndex) -> None:
        state = cb.blockSignals(True)
        cb.setCurrentText(str(index.data(Qt.ItemDataRole.EditRole)))
        cb.blockSignals(state)

    def setModelData(self, cb: QComboBox, model: QAbstractItemModel, index: QModelIndex | QPersistentModelIndex) -> None:
        model.setData(index, cb.currentText(), Qt.ItemDataRole.EditRole)

class FileSectionDelegate(QStyledItemDelegate):

    def createEditor(self, parent: QWidget, option: QStyleOptionViewItem, index: QModelIndex | QPersistentModelIndex) -> QWidget:
        line = QLineEdit(parent)
        self.setEditorData(line, index)
        tex = cast(AmendmentsModel, index.model()).source(index)
        line.setValidator(SectionValidator(tex))
        return line

    def setEditorData(self, line: QLineEdit, index: QModelIndex | QPersistentModelIndex) -> None:
        line.setText(str(index.data(Qt.ItemDataRole.EditRole)))

    def setModelData(self, line: QLineEdit, model: QAbstractItemModel, index: QModelIndex | QPersistentModelIndex) -> None:
        model.setData(index, line.text(), Qt.ItemDataRole.EditRole)

class ProposedAmendmentDelegate(QStyledItemDelegate):

    def createEditor(self, parent: QWidget, option: QStyleOptionViewItem, index: QModelIndex | QPersistentModelIndex) -> QWidget:
        box = QTextEdit(parent)
        self.setEditorData(box, index)
        return box

    def setEditorData(self, box: QTextEdit, index: QModelIndex | QPersistentModelIndex) -> None:
        box.setText(str(index.data(Qt.ItemDataRole.EditRole)))

    def setModelData(self, box: QTextEdit, model: QAbstractItemModel, index: QModelIndex | QPersistentModelIndex) -> None:
        model.setData(index, box.toPlainText(), Qt.ItemDataRole.EditRole)

class AmendmentsModel(QAbstractTableModel):

    sources: dict[str, TeXSource]

    def __init__(self) -> None:
        super().__init__()

        self.amendments: list[list[str]] = []
        self.sources = {}

    def source(self, which: str | QModelIndex | QPersistentModelIndex) -> TeXSource:
        if isinstance(which, str):
            path = which
        else:
            path = self.amendments[which.row()][0]
        if path not in self.sources:
            gh.getTree('skule/bylaws')
            self.sources[path] = TeXSource(gh.getBlob(str(gh.repoPathUrls['skule/bylaws'][path])))
        return self.sources[path]

    def headerData(self, section: int, orientation: Qt.Orientation, role: Qt.ItemDataRole = Qt.ItemDataRole.DisplayRole):
        if role not in ROLES:
            return None
        if orientation == Qt.Orientation.Vertical:
            return None
        try:
            return COLUMNS[section]
        except IndexError:
            return None

    def appendRow(self, path: str | None = None) -> None:
        self.insertRow(self.rowCount())
        if path is not None:
            self.amendments[-1][0] = path
            self.dataChanged.emit(self.index(self.rowCount() - 1, 0),
                                  self.index(self.rowCount() - 1, 0))

    def insertRows(self, row: int, count: int, parent: QModelIndex | QPersistentModelIndex = QModelIndex()) -> bool:
        if row < 0 or row > self.rowCount():
            return False
        self.beginInsertRows(parent, row, row + count - 1)
        self.amendments[row:row] = [[''] * 4 for _ in range(count)]
        self.endInsertRows()
        return True

    def removeRows(self, row: int, count: int, parent: QModelIndex | QPersistentModelIndex = QModelIndex()) -> bool:
        if row + count - 1 < 0 or row > self.rowCount():
            return False
        first = max(0, row)
        last = min(self.rowCount() - 1, row + count - 1)
        self.beginRemoveRows(parent, first, last)
        del self.amendments[first:last+1]
        self.endRemoveRows()
        return True

    def rowCount(self, parent=...) -> int:
        return len(self.amendments)

    def columnCount(self, parent=...) -> int:
        return 4

    def data(self, index: QModelIndex | QPersistentModelIndex, role: Qt.ItemDataRole = Qt.ItemDataRole.DisplayRole):
        if not index.isValid():
            return None
        if role not in ROLES:
            return None
        try:
            return self.amendments[index.row()][index.column()]
        except IndexError:
            return None

    def setData(self, index: QModelIndex | QPersistentModelIndex, value, role: Qt.ItemDataRole = Qt.ItemDataRole.DisplayRole) -> bool:
        if role not in ROLES:
            return False
        if index.column() == 2:
            return False
        try:
            self.amendments[index.row()][index.column()] = value
        except IndexError:
            return False
        if index.column() == 0:
            self.amendments[index.row()][1:] = [''] * 3
            self.dataChanged.emit(self.index(index.row(), 0), self.index(index.row(), 3))
        elif index.column() == 1:
            if not value:
                self.amendments[index.row()][2:] = [''] * 2
            else:
                tex = self.source(index)
                line = tex.lines[tex.linenos[tex.sectionToTuple(value)]]
                line = re.sub(r'^[\s&]*', '', line)
                self.amendments[index.row()][2:] = [line, line]
            self.dataChanged.emit(self.index(index.row(), 1), self.index(index.row(), 3))
        else:
            self.dataChanged.emit(self.index(index.row(), 3), self.index(index.row(), 3))
        return True

    def naturalSort(self) -> None:
        self.amendments.sort(
            key=lambda i: (i[0], i[1] and self.source(i[0]).sectionToTuple(i[1])))
        self.dataChanged.emit(
            self.index(0, 0),
            self.index(self.rowCount() - 1, self.columnCount() - 1),
            list(ROLES)
        )

    def flags(self, index: QModelIndex | QPersistentModelIndex) -> Qt.ItemFlag:
        flag = Qt.ItemFlag.ItemIsEnabled
        if index.column() != 2:
            flag |= Qt.ItemFlag.ItemIsEditable | Qt.ItemFlag.ItemIsSelectable
        return flag

    def open(self, path: str) -> None:
        with open(path, 'r') as f:
            amendments = json.load(f)
        delta = len(amendments) - len(self.amendments)
        if delta > 0:
            self.insertRows(0, delta)
        else:
            self.removeRows(0, -delta)
        self.amendments = amendments
        self.dataChanged.emit(
            self.index(0, 0),
            self.index(self.rowCount() - 1, self.columnCount() - 1),
            list(ROLES)
        )

    def save(self, path: str) -> None:
        with open(path, 'w') as f:
            json.dump(self.amendments, f, indent='\t')

    def exportDocx(self, path: str) -> None:
        document = docx.Document()
        table = document.add_table(self.rowCount() + 1, 4)
        table.autofit = False
        widths = [
            docx.shared.Length(1000000),
            docx.shared.Length(685800),
            docx.shared.Length(1900300),
            docx.shared.Length(1900300),
        ]
        table.style = 'Table Grid'
        for column, cell, text, width in zip(table.columns, table.rows[0].cells, COLUMNS, widths):
            cell.text = text
            cell.width = width
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        for i, row in enumerate(self.amendments, start=1):
            for cell, text, width in zip(table.rows[i].cells, row, widths):
                cell.text = text
                cell.width = width
        document.save(path)
